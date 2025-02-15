import pytesseract
import cv2
import os
import re
import numpy as np
from docx import Document
from docx.shared import RGBColor

# ✅ 指定 Windows 本地 Tesseract-OCR 的路径
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ✅ 设置本地截图文件夹
screenshot_folder = r"screenshots-englishversion"

# ✅ 读取所有截图（只取前 4 张）
image_files = sorted([f for f in os.listdir(screenshot_folder) if f.endswith(".png")])[:4]

# ✅ 创建 Word 文档
doc = Document()
doc.add_heading('题目 + 正确答案', level=1)

def preprocess_image(image_path):
    """ 预处理图像，提高 OCR 识别率 """
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)  # 转灰度
    gray = cv2.GaussianBlur(gray, (3, 3), 0)  # 高斯模糊去噪
    gray = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                 cv2.THRESH_BINARY, 11, 2)  # 二值化
    return gray

def extract_text(image_path):
    """ 使用 OCR 提取文本 """
    processed_img = preprocess_image(image_path)
    text = pytesseract.image_to_string(processed_img, lang="eng+chi_sim")
    return text

def detect_green_areas(image_path):
    """ 识别绿色框选的正确答案 """
    image = cv2.imread(image_path)
    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)

    # 绿色的HSV范围
    lower_green = np.array([35, 100, 100])
    upper_green = np.array([85, 255, 255])

    # 生成遮罩（mask）
    mask = cv2.inRange(hsv, lower_green, upper_green)

    # 识别绿色区域
    contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    green_boxes = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        green_boxes.append((x, y, w, h))  # 记录绿色框的位置

    return green_boxes

for img_file in image_files:
    img_path = os.path.join(screenshot_folder, img_file)
    print(f"📌 处理中: {img_path}")

    # OCR 识别
    text = extract_text(img_path)

    # ✅ 打印 OCR 结果，检查是否正常
    print(f"📜 OCR 识别结果 ({img_file}):\n{text}\n{'-'*80}\n")

    # 修正 OCR 误识别
    text = text.replace("$3", "S3")  # 修正 S3 误识别
    text = text.replace("[AM", "IAM")  # 修正 IAM 误识别
    text = text.replace("0.", "D.")  # 修正 D 选项误识别
    text = text.replace("&", "B.")  # 修正 B 选项误识别

    # 解析 OCR 结果，提取题目和选项
    lines = text.split("\n")
    question = ""
    choices = {}
    correct_answers = []

    option_positions = {}  # 存储选项的 y 坐标

    y_position = 0
    for line in lines:
        line = line.strip()
        if re.match(r"^Q\d+", line):  # 题目编号匹配（如 "Q10"）
            question = line
        elif re.match(r"^[A-E]\.", line):  # 选项匹配（如 "A. xxx"）
            key = line[0]  # 选项 A/B/C/D/E
            choices[key] = line
            option_positions[key] = y_position  # 存储 y 坐标
        y_position += 30  # 估算 y 坐标

    # **检测正确答案（绿色检测）**
    green_boxes = detect_green_areas(img_path)

    # ✅ 打印绿色区域，检查是否正确
    print(f"🟢 绿色区域检测 ({img_file}): {green_boxes}")

    if green_boxes:
        for key, option_text in choices.items():
            option_y = option_positions[key]  # 获取选项的 y 坐标

            # 计算选项与绿色框的最小距离
            for x, y, w, h in green_boxes:
                distance = abs(y - option_y)
                if distance < 50:  # 允许一定误差
                    correct_answers.append(option_text)

    # ✅ 确保每道题都存入 Word 文档
    if question and correct_answers:
        doc.add_paragraph(f"{question}")  # 题目
        for correct in correct_answers:
            run = doc.add_paragraph().add_run(correct)  # 正确答案
            run.font.color.rgb = RGBColor(0, 176, 80)  # 绿色高亮
    else:
        print(f"❌ 错误: {img_file} 没有正确识别到题目或答案！")

# ✅ 保存 Word 文档
word_filename = os.path.join(screenshot_folder, "题目_正确答案_5题测试版.docx")
doc.save(word_filename)
print(f"🎉 识别完成，文档已保存为 {word_filename}")
