#mount google drive
'''
from google.colab import drive
drive.mount('/content/drive')
'''

# in colaboratory 
'''
!apt-get install -y tesseract-ocr
!pip install pytesseract
!pip install opencv-python
!pip install pillow
!pip install python-docx
'''



import pytesseract
import cv2
import os
import re
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import RGBColor
from pytesseract import Output

# 指定 Tesseract OCR 在 Colab 中的路径
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# Google Drive 里存放截图的文件夹
image_folder = "/content/drive/MyDrive/screenshots-englishversion"

# 创建 Word 文档
doc = Document()

def detect_green_regions(image_rgb):
    """检测图像中的绿色区域并返回其Y坐标中心点"""
    hsv = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2HSV)
    lower_green = np.array([35, 50, 50])
    upper_green = np.array([85, 255, 255])
    mask = cv2.inRange(hsv, lower_green, upper_green)
    contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    y_centers = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w * h > 100:  # 过滤小噪点
            y_centers.append(y + h//2)
    return y_centers

def process_image(image_path):
    """处理单张图片并返回结构化数据"""
    # 读取图片
    image = cv2.imread(image_path)
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    
    # 检测绿色区域
    green_y_centers = detect_green_regions(image_rgb)
    
    # 图像预处理
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
    
    # OCR识别并获取文本位置信息
    custom_config = r'--oem 3 --psm 6'
    data = pytesseract.image_to_data(binary, config=custom_config, output_type=Output.DICT)
    
    # 解析文本块
    question_lines = []
    options = []
    current_block = []
    
    for i in range(len(data['text'])):
        text = data['text'][i].strip()
        if not text:
            continue
        
        # 获取坐标信息
        y = data['top'][i]
        
        # 检测题目行（Q开头）
        if re.match(r'^Q\d+', text):
            if current_block:
                question_lines.append(' '.join(current_block))
                current_block = []
            current_block.append(text)
        # 检测选项行（A. B. C. D.）
        elif re.match(r'^[A-D][\.\)]', text):
            # 修正OCR错误
            text = re.sub(r'^([A-D])$', r'\1.', text)  # A → A.
            text = re.sub(r'^0\.', 'D.', text)         # 0. → D.
            options.append({'text': text, 'y': y})
        # 其他文本视为题目内容
        else:
            if current_block:
                current_block.append(text)
    
    if current_block:
        question_lines.append(' '.join(current_block))
    
    # 匹配绿色区域与选项
    correct_answers = []
    for y_center in green_y_centers:
        min_distance = float('inf')
        closest_option = None
        for opt in options:
            distance = abs(opt['y'] - y_center)
            if distance < min_distance:
                min_distance = distance
                closest_option = opt
        if closest_option and min_distance < 50:  # 阈值50像素
            correct_answers.append(closest_option['text'])
    
    return {
        'question': '\n'.join(question_lines),
        'options': [opt['text'] for opt in options],
        'answers': list(set(correct_answers))  # 去重
    }

# 处理所有图片
for idx, image_file in enumerate(sorted(os.listdir(image_folder))):
    if not image_file.lower().endswith(('.png', '.jpg', '.jpeg')):
        continue
    
    image_path = os.path.join(image_folder, image_file)
    result = process_image(image_path)
    
    # 写入Word文档
    if result['question']:
        doc.add_paragraph(f"📌 {result['question']}")
        for opt in result['options']:
            doc.add_paragraph(opt)
        
        if result['answers']:
            p = doc.add_paragraph("✅ 正确答案: ")
            for ans in result['answers']:
                run = p.add_run(ans + " ")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        else:
            doc.add_paragraph("⚠️ 未检测到正确答案", style='IntenseQuote')
        
        doc.add_paragraph("\n" + "-"*50 + "\n")
    
    print(f"处理完成: {image_file} ({idx+1}/{len(os.listdir(image_folder))})")

# 保存文档
output_path = os.path.join(image_folder, "Extracted_Questions.docx")
doc.save(output_path)
print(f"处理完成！文档已保存至：{output_path}")